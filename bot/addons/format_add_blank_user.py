from bot.database.db import db, save_anketa
from bot.addons.send_email import email_scripts
from docx import Document
from docx.shared import Pt
import os


def create_word_file(result_anketa):
    # Создаем папку куда будут сохраняться документы о пользователях.
    os.makedirs("userfile", exist_ok=True)
    # Создаем название документа из имени пользователя и должности.
    name_user = result_anketa["reg_info"]["user_name"].strip()
    position_user = (result_anketa["vacan"].replace("/", "_")).strip()
    name_file = f"{name_user}-{position_user}.docx"
    basedir = os.path.join("userfile", name_file)
    result_anketa["reg_info"]["company"] = result_anketa["company"]
    result_anketa["reg_info"]["vacan"] = result_anketa["vacan"]
    # Делаем словарь с линиями для заполнения
    # Ключи должны быть такие же как и в получаемом файле с информацией об анкетировании
    dict_user_info = {
        "user_name": "ФИО",
        "birth_date": "Дата рождения",
        "location": "Город",
        "relocation": "Переезд",
        "format_job": "Формат работы",
        "salary": "Мин-макс зарплата",
        "number_phone": "Номер телефона",
        "company": "Компания",
        "vacan": "Вакансия",
    }
    # Создаем файл
    # Выбираем шрифт и размер шрифта.
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    # Информация о человеке
    doc.add_heading("Информация о человеке", 1)
    # Будем перебирать dict_user_info по ключу
    # И так как ключи одинаковые с теми данными что мы получаем.
    # Будем вставлять один и тот же ключ в оба словаря.
    for key in dict_user_info:
        doc.add_paragraph(
            f'{dict_user_info[key]}: {result_anketa["reg_info"].get(key, "-")}'
        )
    # Результаты опроса
    doc.add_heading("Результаты опроса", 1)
    for num in range(1, len(result_anketa["answer"]) + 1):
        doc.add_paragraph(f"Вопрос {num}:\n{result_anketa['question'][str(num)]}")
        doc.add_paragraph(f"Ответ {num}:\n{result_anketa['answer'][str(num)]}")
    # Сохраняем файл
    doc.save(basedir)
    # Добавим имя создаваемого файла в словарь отправляемый в базу
    result_anketa["userfile"] = name_file
    save_anketa(db, result_anketa)
    email_scripts(result_anketa)
